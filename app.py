import streamlit as st
import chromadb
from sentence_transformers import SentenceTransformer
from transformers import pipeline
import json
import re
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Page configuration
st.set_page_config(
    page_title="RAG Resume Generator",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: inherit;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-header {
        font-size: 1.5rem;
        font-weight: bold;
        color: inherit;
        margin-bottom: 1rem;
    }
    .resume-section {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
        border-left: 4px solid #1f77b4;
        color: #333;
    }
    .skill-tag {
        background-color: #e3f2fd;
        color: #1976d2;
        padding: 0.25rem 0.5rem;
        border-radius: 0.25rem;
        margin: 0.25rem;
        display: inline-block;
        font-size: 0.875rem;
    }
    /* Ensure text is readable in both light and dark modes */
    .stTextArea textarea {
        color: #333 !important;
    }
    .stMarkdown {
        color: inherit;
    }
</style>
""", unsafe_allow_html=True)

@st.cache_resource
def load_models():
    """Load and cache the embedding and generation models"""
    embedder = SentenceTransformer("all-MiniLM-L6-v2")
    generator = pipeline("text2text-generation", model="google/flan-t5-large")
    return embedder, generator

@st.cache_resource
def setup_chromadb():
    """Setup and cache ChromaDB connection"""
    client = chromadb.PersistentClient(path="./data")
    collection = client.get_or_create_collection(name="resume_blocks")
    return collection

def extract_skills_from_query(query):
    """Extract potential skills from the job query"""
    skill_keywords = {
        'python': 'Python', 'javascript': 'JavaScript', 'java': 'Java', 'react': 'React',
        'node.js': 'Node.js', 'aws': 'AWS', 'docker': 'Docker', 'kubernetes': 'Kubernetes',
        'machine learning': 'Machine Learning', 'ml': 'Machine Learning', 'ai': 'AI',
        'data science': 'Data Science', 'sql': 'SQL', 'nosql': 'NoSQL',
        'git': 'Git', 'agile': 'Agile', 'scrum': 'Scrum', 'devops': 'DevOps',
        'frontend': 'Frontend', 'backend': 'Backend', 'full-stack': 'Full-Stack',
        'cloud': 'Cloud Computing', 'microservices': 'Microservices'
    }
    
    found_skills = []
    query_lower = query.lower()
    for keyword, skill in skill_keywords.items():
        if keyword in query_lower:
            found_skills.append(skill)
    
    return found_skills

def create_pdf_resume(resume_data):
    """Create a PDF resume"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=12,
        textColor=colors.darkblue
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=8,
        textColor=colors.darkblue
    )
    normal_style = styles['Normal']
    
    # Add title
    story.append(Paragraph("Professional Resume", title_style))
    story.append(Spacer(1, 12))
    
    # Add sections
    sections = [
        ("Professional Summary", resume_data.get("summary", "")),
        ("Work Experience", resume_data.get("experience", "")),
        ("Technical Skills", resume_data.get("skills", "")),
        ("Key Projects", resume_data.get("projects", "")),
        ("Education", resume_data.get("education", ""))
    ]
    
    for section_title, content in sections:
        if content:
            story.append(Paragraph(section_title, heading_style))
            story.append(Paragraph(content, normal_style))
            story.append(Spacer(1, 12))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_docx_resume(resume_data):
    """Create a DOCX resume"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Professional Resume', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add sections
    sections = [
        ("Professional Summary", resume_data.get("summary", "")),
        ("Work Experience", resume_data.get("experience", "")),
        ("Technical Skills", resume_data.get("skills", "")),
        ("Key Projects", resume_data.get("projects", "")),
        ("Education", resume_data.get("education", ""))
    ]
    
    for section_title, content in sections:
        if content:
            doc.add_heading(section_title, level=1)
            doc.add_paragraph(content)
            doc.add_paragraph()  # Add spacing
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def remove_duplicates(text):
    """Remove duplicate sentences and clean up the text"""
    if not text:
        return text
    
    # Split by sentences and remove duplicates
    sentences = text.split('. ')
    seen = set()
    unique_sentences = []
    
    for sentence in sentences:
        sentence = sentence.strip()
        if sentence and sentence not in seen:
            unique_sentences.append(sentence)
            seen.add(sentence)
    
    # Join back with proper punctuation
    result = '. '.join(unique_sentences)
    if result and not result.endswith('.'):
        result += '.'
    
    return result

def generate_resume_section(job_role, context, section_type="summary", generator=None):
    """Generate different types of resume sections"""
    
    prompts = {
        "summary": f"""Write a concise, professional summary for a resume targeting the job role: {job_role}. 
        Use the context below to create a compelling summary that highlights relevant skills and experience. 
        Make it sound professional and avoid repetition. Context: {context}""",
        
        "experience": f"""Create 3-4 unique bullet points describing relevant work experience for a {job_role} role. 
        Each bullet should be specific and highlight different achievements or responsibilities. 
        Use the context below and avoid repeating phrases. Context: {context}""",
        
        "skills": f"""List technical skills and competencies relevant for a {job_role} position. 
        Organize them by category (e.g., Programming Languages, Frameworks, Tools). 
        Use the context below and avoid listing the same skill multiple times. Context: {context}""",
        
        "projects": f"""Describe 2-3 unique projects relevant for a {job_role} role. 
        Each project should have a clear description of the technology used and outcomes achieved. 
        Use the context below and avoid repetition. Context: {context}""",
        
        "education": f"""Write an education section appropriate for a {job_role} position. 
        Include relevant degrees, certifications, and any specialized training. 
        Use the context below and avoid repetition. Context: {context}"""
    }
    
    if generator is None:
        _, generator = load_models()
    
    prompt = prompts.get(section_type, prompts["summary"])
    
    # Generate with better parameters
    response = generator(
        prompt, 
        max_length=400, 
        do_sample=True, 
        temperature=0.8,
        top_p=0.9,
        repetition_penalty=1.2
    )[0]["generated_text"]
    
    # Clean up the response
    cleaned_response = remove_duplicates(response)
    return cleaned_response

def main():
    # Header
    st.markdown('<h1 class="main-header">💼 RAG Resume Generator</h1>', unsafe_allow_html=True)
    
    # Load models and setup
    embedder, generator = load_models()
    collection = setup_chromadb()
    
    # Sidebar for configuration
    with st.sidebar:
        st.markdown("### ⚙️ Configuration")
        
        # Job role input
        job_role = st.text_input(
            "🎯 Job Role/Description",
            placeholder="e.g., Senior Python Developer with ML experience",
            help="Describe the job role you're targeting"
        )
        
        # Experience level
        experience_level = st.selectbox(
            "📊 Experience Level",
            ["Entry Level", "Mid Level", "Senior Level", "Lead/Manager"],
            help="Select your experience level"
        )
        
        # Industry focus
        industry = st.selectbox(
            "🏢 Industry Focus",
            ["Technology", "Finance", "Healthcare", "E-commerce", "Education", "General"],
            help="Select your target industry"
        )
        
        # Generate button
        generate_btn = st.button("🚀 Generate Resume", type="primary")
    
    # Main content area
    if job_role and generate_btn:
        st.markdown("---")
        
        # Extract skills from query
        detected_skills = extract_skills_from_query(job_role)
        
        # Show detected skills
        if detected_skills:
            st.markdown("### 🎯 Detected Skills")
            for skill in detected_skills:
                st.markdown(f'<span class="skill-tag">{skill}</span>', unsafe_allow_html=True)
        
        # Query the RAG system
        query_embedding = embedder.encode(job_role)
        results = collection.query(
            query_embeddings=[query_embedding], 
            n_results=8
        )
        
        # Combine context
        context = " ".join([doc for result in results["documents"] for doc in result])
        
        # Generate different resume sections
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown('<h3 class="section-header">📋 Professional Summary</h3>', unsafe_allow_html=True)
            summary = generate_resume_section(job_role, context, "summary", generator)
            st.markdown(f'<div class="resume-section">{summary}</div>', unsafe_allow_html=True)
            
            st.markdown('<h3 class="section-header">💼 Work Experience</h3>', unsafe_allow_html=True)
            experience = generate_resume_section(job_role, context, "experience", generator)
            st.markdown(f'<div class="resume-section">{experience}</div>', unsafe_allow_html=True)
            
            st.markdown('<h3 class="section-header">🎓 Education</h3>', unsafe_allow_html=True)
            education = generate_resume_section(job_role, context, "education", generator)
            st.markdown(f'<div class="resume-section">{education}</div>', unsafe_allow_html=True)
        
        with col2:
            st.markdown('<h3 class="section-header">🛠️ Technical Skills</h3>', unsafe_allow_html=True)
            skills = generate_resume_section(job_role, context, "skills", generator)
            st.markdown(f'<div class="resume-section">{skills}</div>', unsafe_allow_html=True)
            
            st.markdown('<h3 class="section-header">🚀 Key Projects</h3>', unsafe_allow_html=True)
            projects = generate_resume_section(job_role, context, "projects", generator)
            st.markdown(f'<div class="resume-section">{projects}</div>', unsafe_allow_html=True)
        
        # Full resume download
        st.markdown("---")
        st.markdown('<h3 class="section-header">📄 Complete Resume</h3>', unsafe_allow_html=True)
        
        full_resume = f"""
# Professional Resume

## Professional Summary
{summary}

## Work Experience
{experience}

## Technical Skills
{skills}

## Key Projects
{projects}

## Education
{education}
        """
        
        st.text_area("Complete Resume", value=full_resume, height=400)
        
        # Download options
        st.markdown("### 📥 Download Options")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # TXT Download
            st.download_button(
                label="📄 Download as TXT",
                data=full_resume,
                file_name=f"resume_{job_role.replace(' ', '_').lower()}.txt",
                mime="text/plain",
                help="Download as plain text file"
            )
        
        with col2:
            # PDF Download
            resume_data = {
                "summary": summary,
                "experience": experience,
                "skills": skills,
                "projects": projects,
                "education": education
            }
            pdf_buffer = create_pdf_resume(resume_data)
            st.download_button(
                label="📋 Download as PDF",
                data=pdf_buffer.getvalue(),
                file_name=f"resume_{job_role.replace(' ', '_').lower()}.pdf",
                mime="application/pdf",
                help="Download as PDF document"
            )
        
        with col3:
            # DOCX Download
            docx_buffer = create_docx_resume(resume_data)
            st.download_button(
                label="📝 Download as DOCX",
                data=docx_buffer.getvalue(),
                file_name=f"resume_{job_role.replace(' ', '_').lower()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Download as Word document"
            )
        
        # RAG Context Information
        with st.expander("🔍 RAG Context Information"):
            st.markdown("**Retrieved Context:**")
            st.text(context[:500] + "..." if len(context) > 500 else context)
            
            st.markdown("**Query:**")
            st.text(job_role)
            
            st.markdown("**Number of retrieved documents:**")
            st.text(len(results["documents"][0]) if results["documents"] else 0)
    
    elif not job_role:
        # Welcome message
        st.markdown("""
        ## 🎯 How to Use This RAG Resume Generator
        
        This application uses **Retrieval-Augmented Generation (RAG)** to create personalized resumes:
        
        ### 🔍 **RAG Process:**
        1. **Retrieval**: Searches through a database of resume content based on your job description
        2. **Generation**: Uses AI to create relevant resume sections using the retrieved context
        3. **Personalization**: Tailors content to your specific job role and requirements
        
        ### 📝 **To Get Started:**
        1. Enter your target job role in the sidebar
        2. Select your experience level and industry focus
        3. Click "Generate Resume" to create your personalized resume
        
        ### 🎨 **Features:**
        - **Smart Skill Detection**: Automatically identifies relevant skills from your job description
        - **Multiple Sections**: Generates summary, experience, skills, projects, and education
        - **Industry Focus**: Tailors content based on your target industry
        - **Download Ready**: Export your generated resume as a text file
        
        ### 💡 **Tips for Better Results:**
        - Be specific about the job role (e.g., "Senior Python Developer with ML experience")
        - Include key technologies or skills in your description
        - Mention the industry or company type you're targeting
        """)
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #666;'>
        <p>💼 RAG Resume Generator | Powered by ChromaDB + Sentence Transformers + T5</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()